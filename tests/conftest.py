"""Shared pytest fixtures.

Provides:

- ``tmp_storage``  -- per-test scratch directory wired into a fresh
  ``SQLiteStorage`` + ``VectorStore``.
- ``sample_docx``  -- bytes of a programmatically generated DOCX.
- ``sample_pdf``   -- bytes of a programmatically generated PDF (with real
  extractable text via pypdf + a Type1 font).
- ``chroma_settings`` -- the chromadb Settings object used in tests.
- ``mock_embedding_client`` -- a stand-in for ``EmbeddingClient`` that
  returns deterministic vectors (no network).

LLM mocking is done by patching ``app.llm.minimax_client.AsyncOpenAI``
where needed; embedding clients use the local fake so we never touch
the network.
"""

from __future__ import annotations

import asyncio
import io
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional

import pytest
import yaml
from chromadb.config import Settings as ChromaSettings

from app.config import (
    AuthConfig,
    ServerConfig,
    ChatConfig,
    EmbeddingConfig,
    RagConfig,
    Settings,
    StorageConfig,
    UploadConfig,
    reset_settings_cache,
)
from app.kb.storage import SQLiteStorage
from app.rag.vectorstore import VectorStore


# ---------------------------------------------------------------------------
# Settings / directories
# ---------------------------------------------------------------------------


@pytest.fixture
def tmp_storage(tmp_path: Path) -> Dict[str, Path]:
    """Return a dict of tmp paths for storage / chroma / uploads."""
    root = tmp_path
    upload_dir = root / "uploads"
    chroma_dir = root / "chroma"
    meta_db = root / "meta.sqlite3"
    upload_dir.mkdir(parents=True, exist_ok=True)
    chroma_dir.mkdir(parents=True, exist_ok=True)
    return {
        "root": root,
        "upload_dir": upload_dir,
        "chroma_dir": chroma_dir,
        "meta_db": meta_db,
    }


@pytest.fixture
def settings(tmp_storage: Dict[str, Path]) -> Settings:
    """Build a Settings instance pointed at the per-test tmp paths."""
    s = Settings(
        project_root=tmp_storage["root"],
        server=ServerConfig(),
        chat=ChatConfig(),
        embedding=EmbeddingConfig(),
        storage=StorageConfig(
            upload_dir=str(tmp_storage["upload_dir"]),
            chroma_dir=str(tmp_storage["chroma_dir"]),
            meta_db=str(tmp_storage["meta_db"]),
        ),
        rag=RagConfig(),
        upload=UploadConfig(),
        auth=AuthConfig(
            jwt_secret="test-jwt-secret-do-not-use-in-prod",
            admin_username="root",
            admin_password="rootpw",
            bcrypt_rounds=4,
            cookie_secure=False,
        ),
    )
    # Best-effort: if the global ``get_settings`` cache was monkey-patched
    # by another fixture, the cache_clear call would raise. Swallow that
    # -- the per-test ``settings`` fixture always returns a fresh instance
    # anyway, so the cached singleton is irrelevant here.
    try:
        reset_settings_cache()
    except AttributeError:
        pass
    return s


@pytest.fixture
def sqlite_storage(settings: Settings) -> SQLiteStorage:
    return SQLiteStorage(settings)


@pytest.fixture
def vectorstore(settings: Settings) -> VectorStore:
    return VectorStore(settings)


# ---------------------------------------------------------------------------
# Sample-file generation
# ---------------------------------------------------------------------------


@pytest.fixture(scope="session")
def sample_text() -> str:
    """Mixed Chinese + English text used to build the .txt/.md fixtures."""
    return (
        "youfu-known 是一个本地化的个人知识库 + RAG 系统。\n"
        "它支持上传 PDF / Word / Markdown / HTML / 纯文本, "
        "然后基于 Qwen3-Embedding + Chroma 向量库做语义检索, "
        "最后调用 MiniMax 大模型生成带引用的回答。\n\n"
        "Key features:\n"
        "- 完全本地存储文件与向量元数据, 只在调用 LLM / Embedding 时走外网。\n"
        "- 支持多知识库: 每个知识库 = 一个独立的 Chroma Collection。\n"
        "- 引用透明: 每个回答会带上 [n] 形式的来源标注。\n\n"
        "Example workflow:\n"
        "1. 用户在前端拖拽上传一份 PDF;\n"
        "2. 后端把文件落到 storage/uploads/{kb_id}/{doc_id}.pdf;\n"
        "3. 后台任务加载 -> 切块 -> Embedding -> 写入 Chroma;\n"
        "4. 状态机: pending -> processing -> ready / failed;\n"
        "5. 用户提问时, 先召回 top-k 个片段, 再交给 MiniMax 汇总成答案。"
    )


@pytest.fixture(scope="session")
def samples_dir() -> Path:
    """Return the path to tests/samples (created on demand)."""
    return Path(__file__).resolve().parent / "samples"


def _write_text(path: Path, body: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(body, encoding="utf-8")
    return path


@pytest.fixture(scope="session")
def sample_txt(sample_text: str, samples_dir: Path) -> Path:
    return _write_text(samples_dir / "a.txt", sample_text)


@pytest.fixture(scope="session")
def sample_md(samples_dir: Path) -> Path:
    body = (
        "# youfu-known 入门\n\n"
        "这是第一个示例 Markdown 文件, 用于测试 Markdown 解析器。\n\n"
        "## 主要特性\n\n"
        "- 完全本地化\n"
        "- 支持多种文档格式\n"
        "- 基于 RAG 的问答\n\n"
        "## 代码块示例\n\n"
        "```python\n"
        "from app.config import get_settings\n"
        'print(get_settings().chat.model)\n'
        "```\n\n"
        "## 表格示例\n\n"
        "| 名称 | 描述 |\n"
        "| --- | --- |\n"
        "| LLM | MiniMax |\n"
        "| Embedding | DashScope Qwen3 |\n"
    )
    return _write_text(samples_dir / "a.md", body)


@pytest.fixture(scope="session")
def sample_html(samples_dir: Path) -> Path:
    body = (
        "<!doctype html>\n"
        "<html><head><meta charset=\"utf-8\"><title>sample</title></head>\n"
        "<body>\n"
        "<h1>Sample HTML</h1>\n"
        "<p>这是 youfu-known 解析器测试用的简单 HTML 文件.</p>\n"
        "<ul>\n"
        "<li>列表项 A</li>\n"
        "<li>列表项 B</li>\n"
        "</ul>\n"
        "<script>alert('should be stripped');</script>\n"
        "</body></html>\n"
    )
    return _write_text(samples_dir / "a.html", body)


@pytest.fixture(scope="session")
def sample_docx(samples_dir: Path) -> Path:
    """Generate a real DOCX with python-docx on first call."""
    from docx import Document

    p = samples_dir / "a.docx"
    if not p.exists():
        doc = Document()
        doc.add_heading("youfu-known DOCX 测试", level=1)
        doc.add_paragraph(
            "这是一个由 python-docx 程序化生成的 Word 文档, 用于验证 parser_docx.py。"
        )
        doc.add_paragraph(
            "DOCX 解析按段落返回, page 字段留 None。"
        )
        doc.add_paragraph(
            "Project: youfu-known\n"
            "Owner: youfu\n"
            "Tags: rag, chromadb, minimax, dashscope"
        )
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
    return p


@pytest.fixture(scope="session")
def sample_pdf(samples_dir: Path) -> Path:
    """Generate a real PDF with extractable text using pypdf primitives."""
    p = samples_dir / "a.pdf"
    if not p.exists():
        from pypdf import PdfWriter
        from pypdf.generic import (
            DecodedStreamObject,
            DictionaryObject,
            NameObject,
        )

        p.parent.mkdir(parents=True, exist_ok=True)
        writer = PdfWriter()

        # Add 2 pages with different text each.
        for body in (
            "Page one of the sample PDF. youfu-known RAG system.",
            "Page two mentions DashScope embedding and Chroma vector store.",
        ):
            page = writer.add_blank_page(width=612, height=792)

            # Wire up a minimal Type1 font resource so pypdf can decode text.
            font = DictionaryObject()
            font[NameObject("/Type")] = NameObject("/Font")
            font[NameObject("/Subtype")] = NameObject("/Type1")
            font[NameObject("/BaseFont")] = NameObject("/Helvetica")

            fonts = DictionaryObject()
            fonts[NameObject("/F1")] = font

            resources = DictionaryObject()
            resources[NameObject("/Font")] = fonts

            page_obj = writer._pages["/Kids"][-1].get_object()
            page_obj[NameObject("/Resources")] = resources

            safe = body.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
            stream_data = f"BT /F1 14 Tf 72 720 Td ({safe}) Tj ET".encode("latin-1")
            stream = DecodedStreamObject()
            stream.set_data(stream_data)
            writer._objects.append(stream)
            stream_ref = writer._add_object(stream)
            page_obj[NameObject("/Contents")] = stream_ref

        with p.open("wb") as fh:
            writer.write(fh)
    return p


# ---------------------------------------------------------------------------
# Sample directory fixture (pre-generates all samples)
# ---------------------------------------------------------------------------


@pytest.fixture(scope="session")
def all_samples(
    samples_dir: Path,
    sample_txt: Path,
    sample_md: Path,
    sample_html: Path,
    sample_docx: Path,
    sample_pdf: Path,
) -> Dict[str, Path]:
    return {
        "txt": sample_txt,
        "md": sample_md,
        "html": sample_html,
        "docx": sample_docx,
        "pdf": sample_pdf,
    }


# ---------------------------------------------------------------------------
# Fake embedding / chat clients (no network)
# ---------------------------------------------------------------------------


@dataclass
class FakeEmbeddingClient:
    """Deterministic, in-memory EmbeddingClient used in tests."""

    dim: int = 16
    batch_size: int = 25
    model: str = "fake-embedding-v1"

    async def aembed(self, texts: List[str]) -> List[List[float]]:
        # Hash each text into a fixed-length vector. Same text -> same vector.
        import hashlib

        out: List[List[float]] = []
        for text in texts:
            digest = hashlib.sha256(text.encode("utf-8")).digest()
            vec = []
            for i in range(self.dim):
                byte = digest[i % len(digest)]
                vec.append(((byte / 255.0) - 0.5) * 2.0)  # [-1, 1]
            out.append(vec)
        return out

    async def aembed_iter(self, texts: Iterable[str]) -> List[List[float]]:
        return await self.aembed(list(texts))


@dataclass
class FakeChatClient:
    """Deterministic chat client that echoes a canned reply."""

    reply: str = "这是来自 fake LLM 的固定回答 [1]."

    async def achat(self, messages, **kw) -> str:  # noqa: ARG002
        return self.reply

    async def astream(self, messages, **kw):  # noqa: ARG002
        for ch in self.reply:
            yield ch


@pytest.fixture
def fake_embedding_client() -> FakeEmbeddingClient:
    return FakeEmbeddingClient()


@pytest.fixture
def fake_chat_client() -> FakeChatClient:
    return FakeChatClient()


@pytest.fixture
def kb_service(
    settings,
    sqlite_storage: SQLiteStorage,
    vectorstore: VectorStore,
    fake_embedding_client,
):
    """Default KBService wired up with the fake embedding client."""
    from app.kb.service import KBService
    from app.rag.embedder import Embedder

    return KBService(
        storage=sqlite_storage,
        vectorstore=vectorstore,
        embedder=Embedder(fake_embedding_client),
        settings=settings,
    )


# ---------------------------------------------------------------------------
# TestClient harness (shared by tests that need to hit the HTTP layer)
# ---------------------------------------------------------------------------


def _make_settings_for_testclient(project_root: Path, tmp: Path) -> Settings:
    """Build a Settings instance pointing every storage path at ``tmp``.

    ``project_root`` becomes the project root (used by config resolver
    helpers), but all actual data dirs are redirected to ``tmp``.
    """
    cfg_path = project_root / "config.yaml"
    with cfg_path.open() as f:
        cfg = yaml.safe_load(f)

    cfg["storage"]["upload_dir"] = str(tmp / "uploads")
    cfg["storage"]["chroma_dir"] = str(tmp / "chroma")
    cfg["storage"]["meta_db"] = str(tmp / "meta.sqlite3")
    # Inject placeholder API keys so pydantic Settings validate.
    cfg["chat"]["api_key"] = "test-chat-key"
    cfg["embedding"]["api_key"] = "test-embed-key"
    return Settings(project_root=project_root, **cfg)


@pytest.fixture()
def api_settings(
    monkeypatch: pytest.MonkeyPatch, tmp_path: Path
) -> Settings:
    """Settings object wired up for the TestClient harness.

    Redirects every storage dir at ``tmp_path`` so the app cannot leak
    real data. Also injects dummy API keys so pydantic Settings accept,
    and pre-populates an admin bootstrap so the auth lifespan can seed
    the initial admin user.
    """
    import yaml  # local import keeps top-level deps lean

    project_root = Path(__file__).resolve().parent.parent
    monkeypatch.setenv("YOUFU_KNOWN_ROOT", str(project_root))
    # Pre-populate admin bootstrap creds so the lifespan handler can
    # seed the initial admin when the users table is empty.
    monkeypatch.setenv("YOUFU_ADMIN_USERNAME", "root")
    monkeypatch.setenv("YOUFU_ADMIN_PASSWORD", "rootpw")
    monkeypatch.setenv("YOUFU_COOKIE_SECURE", "false")
    monkeypatch.setenv(
        "YOUFU_JWT_SECRET", "test-jwt-secret-do-not-use-in-prod"
    )
    # Drop any cached singleton so the env-var change takes effect.
    import app.config as config_mod
    config_mod.get_settings.cache_clear()  # type: ignore[attr-defined]

    import importlib
    importlib.reload(config_mod)

    settings = _make_settings_for_testclient(project_root, tmp_path)
    # Lower bcrypt cost so register/login tests stay snappy.
    settings.auth.bcrypt_rounds = 4
    settings.auth.session_hours = 24
    settings.auth.refresh_days = 30
    settings.auth.cookie_secure = False
    settings.auth.jwt_secret = "test-jwt-secret-do-not-use-in-prod"
    settings.auth.admin_username = "root"
    settings.auth.admin_password = "rootpw"
    import app.deps as deps
    monkeypatch.setattr(config_mod, "get_settings", lambda: settings)
    monkeypatch.setattr(deps, "get_settings", lambda: settings)
    return settings


@pytest.fixture()
def client(api_settings: Settings) -> Iterator[TestClient]:
    """Spin up the FastAPI app pointed at the temp storage.

    The lifespan handler runs on context entry, which (because
    ``api_settings`` seeded admin credentials) will bootstrap the
    initial admin. Tests that need to log in can use ``admin_client``.
    """
    from fastapi.testclient import TestClient

    from main import create_app

    app = create_app()
    with TestClient(app, raise_server_exceptions=False) as c:
        yield c


@pytest.fixture()
def admin_client(client: TestClient) -> TestClient:
    """Client with the bootstrapped admin already logged in.

    Most existing tests want a happy-path admin session; this fixture
    removes the boilerplate from each test. If you need a fresh
    session (e.g. to test logout), use ``client`` directly.
    """
    r = client.post(
        "/api/auth/login",
        json={"username": "root", "password": "rootpw"},
    )
    assert r.status_code == 200, r.text
    return client


class _FakeRetriever:
    """Drop-in replacement for Retriever used by the chat endpoint."""

    def __init__(self) -> None:
        from unittest.mock import AsyncMock

        from app.rag.retriever import Citation, RagResult

        self.ask = AsyncMock(
            return_value=RagResult(
                answer=(
                    "根据资料, MiniMax Embedding 接口地址为 "
                    "https://api.MiniMax.chat/v1/embeddings [1]。"
                ),
                citations=[
                    Citation(
                        n=1,
                        doc_id="doc-abc",
                        doc_filename="minimax_docs.md",
                        chunk_idx=3,
                        chunk_id="doc-abc::3",
                        score=0.82,
                        text=(
                            "MiniMax Embedding 接口地址为 "
                            "https://api.MiniMax.chat/v1/embeddings ..."
                        ),
                    )
                ],
            )
        )


@pytest.fixture()
def mock_retriever(client: TestClient) -> _FakeRetriever:
    """Swap the lifespan-built retriever with a mock."""
    fake = _FakeRetriever()
    client.app.state.retriever = fake  # type: ignore[attr-defined]
    return fake


@pytest.fixture()
def mock_embedder(client: TestClient) -> None:
    """Replace the embedding client with one that returns random vectors.

    The Chroma collection persists in tmp_path, so upserting fake
    embeddings is fine -- the test only asserts the pipeline reaches
    READY status, not that retrievals are sensible.
    """
    import random

    class _FakeEmbedClient:
        dim = 1024  # must match real DashScope embedding dim
        model = "fake-embed"

        async def aembed(self, texts):
            return [[random.random() for _ in range(self.dim)] for _ in texts]

    fake = _FakeEmbedClient()
    client.app.state.embed_client = fake  # type: ignore[attr-defined]
    # KBService holds its own reference to the embedder; swap its inner
    # client too so ingest_document uses the fake.
    embedder = client.app.state.embedder  # type: ignore[attr-defined]
    embedder._client = fake  # type: ignore[attr-defined]


# ---------------------------------------------------------------------------
# Async helper
# ---------------------------------------------------------------------------


def run_async(coro):
    """Run a coroutine in a fresh event loop (sync test convenience)."""
    return asyncio.get_event_loop().run_until_complete(coro) if False else asyncio.run(coro)